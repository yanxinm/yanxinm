#!/usr/bin/env python3
"""
将 Markdown 公文转为 Word (.docx) 格式
遵循 GB/T 9704-2012 党政机关公文格式

用法:
  python3 convert-markdown-to-gov-docx.py <输入.md> [输出.docx]

如不指定输出路径，自动将后缀改为 .docx 保存在同目录。

依赖: python-docx
"""

import re
import sys
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ─── 页面设置 (GB/T 9704-2012) ───
PAGE_CONFIG = {
    "top": 3.7,      # 上 37mm
    "bottom": 3.5,   # 下 35mm
    "left": 2.8,     # 左 28mm
    "right": 2.6,    # 右 26mm
}

# ─── 样式 ───

def set_font(run, name_cn='仿宋', name_en='Times New Roman', size=Pt(16), bold=False):
    run.font.size = size
    run.font.bold = bold
    run.font.name = name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name_cn)


def set_line_spacing(paragraph, spacing=Pt(28)):
    paragraph.paragraph_format.line_spacing = spacing


def add_title(doc, text):
    """公文标题：二号黑体 (18pt) 居中"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    set_line_spacing(p, Pt(28))
    run = p.add_run(text)
    set_font(run, name_cn='黑体', size=Pt(18), bold=True)
    return p


def add_heading1(doc, text):
    """一级标题 一、二、三：三号黑体 (16pt)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    set_line_spacing(p, Pt(28))
    run = p.add_run(text)
    set_font(run, name_cn='黑体', size=Pt(16), bold=True)
    return p


def add_heading2(doc, text):
    """二级标题（一）（二）：三号楷体 (16pt)"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    set_line_spacing(p, Pt(28))
    run = p.add_run(text)
    set_font(run, name_cn='楷体', size=Pt(16), bold=False)
    return p


def add_body(doc, text):
    """正文：三号仿宋 (16pt) 首行缩进2字符"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Pt(32)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    set_line_spacing(p, Pt(28))
    run = p.add_run(text)
    set_font(run, name_cn='仿宋', size=Pt(16))
    return p


def add_bold_body(doc, text):
    """加粗正文段"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Pt(32)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    set_line_spacing(p, Pt(28))
    run = p.add_run(text)
    set_font(run, name_cn='仿宋', size=Pt(16), bold=True)
    return p


def add_mixed_paragraph(doc, segments):
    """混合格式段落：[(text, bold), ...]"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Pt(32)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    set_line_spacing(p, Pt(28))
    for text, bold in segments:
        run = p.add_run(text)
        set_font(run, name_cn='仿宋', size=Pt(16), bold=bold)
    return p


def add_italic_source(doc, text):
    """数据来源：五号仿宋斜体"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.first_line_indent = Pt(32)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    set_line_spacing(p, Pt(20))
    run = p.add_run(text)
    set_font(run, name_cn='仿宋', size=Pt(10.5))
    run.font.italic = True
    return p


def add_code_block(doc, text):
    """代码/公式块：五号仿宋"""
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.first_line_indent = Pt(0)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        set_line_spacing(p, Pt(20))
        run = p.add_run(line)
        set_font(run, name_cn='仿宋', size=Pt(10.5))
        run.font.italic = True


def add_separator(doc):
    """分隔线"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('—' * 30)
    set_font(run, name_cn='仿宋', size=Pt(10))
    return p


def add_table_from_markdown(doc, md_lines):
    """Markdown 表格 → Word 表格"""
    headers = [h.strip() for h in md_lines[0].strip().split('|')[1:-1]]
    data_rows = []
    for line in md_lines[2:]:
        line = line.strip()
        if not line or line.startswith('|---'):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if cells:
            data_rows.append(cells)

    cols = len(headers)
    if cols == 0:
        return None

    table = doc.add_table(rows=len(data_rows) + 1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, name_cn='仿宋', size=Pt(12), bold=True)

    # Data rows
    for r_idx, row_data in enumerate(data_rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            parts = re.split(r'(\*\*.*?\*\*)', cell_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_font(run, name_cn='仿宋', size=Pt(12), bold=True)
                else:
                    run = p.add_run(part)
                    set_font(run, name_cn='仿宋', size=Pt(12))

    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    # Spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)

    return table


# ─── 主流程 ───

def convert(md_path, docx_path=None):
    """将 Markdown 公文转为 Word"""

    if not docx_path:
        docx_path = os.path.splitext(md_path)[0] + '.docx'

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(PAGE_CONFIG["top"])
    section.bottom_margin = Cm(PAGE_CONFIG["bottom"])
    section.left_margin = Cm(PAGE_CONFIG["left"])
    section.right_margin = Cm(PAGE_CONFIG["right"])

    in_table = False
    table_lines = []
    in_code = False
    code_lines = []
    first_line = True

    for raw_line in lines:
        line = raw_line.rstrip()

        # Code blocks
        if line.startswith('```'):
            if in_code:
                if code_lines:
                    add_code_block(doc, '\n'.join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue

        # Tables
        if line.startswith('|') and line.endswith('|'):
            if not in_table:
                in_table = True
                table_lines = [line]
            else:
                table_lines.append(line)
            continue
        else:
            if in_table and table_lines:
                add_table_from_markdown(doc, table_lines)
                in_table = False
                table_lines = []

        # Skip empty
        if not line:
            continue
        if line.startswith('---') and len(line.strip('-')) < 5:
            add_separator(doc)
            continue

        # Title (first # line)
        if first_line and line.startswith('# ') and not line.startswith('## '):
            add_title(doc, re.sub(r'^#\s*', '', line))
            first_line = False
            continue

        # Headings
        if line.startswith('### '):
            add_heading2(doc, re.sub(r'^#+\s*', '', line))
            continue
        if line.startswith('## '):
            add_heading1(doc, re.sub(r'^#+\s*', '', line))
            continue

        # 数据来源 (italic)
        if line.startswith('*数据来源'):
            add_italic_source(doc, line)
            continue

        # Bold-only lines
        if line.startswith('**') and line.endswith('**') and len(line) < 100:
            add_bold_body(doc, line.replace('**', ''))
            continue

        # Mixed bold inline
        if '**' in line:
            parts = re.split(r'(\*\*.*?\*\*)', line)
            segments = []
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    segments.append((part[2:-2], True))
                else:
                    segments.append((part, False))
            add_mixed_paragraph(doc, segments)
            continue

        # Default body
        add_body(doc, line)

    # Flush table
    if in_table and table_lines:
        add_table_from_markdown(doc, table_lines)

    doc.save(docx_path)
    return docx_path


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python3 convert-markdown-to-gov-docx.py <输入.md> [输出.docx]")
        sys.exit(1)

    md_path = sys.argv[1]
    docx_path = sys.argv[2] if len(sys.argv) > 2 else None

    if not os.path.isfile(md_path):
        print(f"❌ 文件不存在: {md_path}")
        sys.exit(1)

    result = convert(md_path, docx_path)
    print(f"✅ 已生成: {result}")
    print(f"   输入: {md_path}")
    print(f"   排版: GB/T 9704-2012 (上3.7 下3.5 左2.8 右2.6cm)")
